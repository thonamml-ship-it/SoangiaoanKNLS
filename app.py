import streamlit as st
import google.generativeai as genai
from PIL import Image
import tempfile
import os
import io
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- 1. CẤU HÌNH TRANG ---
st.set_page_config(page_title="Trợ lý Giáo án NLS", page_icon="📘", layout="centered")

FILE_KHUNG_NANG_LUC = "khungnanglucso.pdf"

# --- 2. HÀM XỬ LÝ WORD (ĐÃ CẬP NHẬT: GẠCH ĐẦU DÒNG THỦ CÔNG) ---

def add_formatted_text(paragraph, text):
    """Hàm in đậm và ép font Times New Roman"""
    paragraph.style.font.name = 'Times New Roman'
    paragraph.style.font.size = Pt(14)
    
    parts = re.split(r'(\*\*.*?\*\*)', text) 
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            clean = part[2:-2]
            run = paragraph.add_run(clean)
            run.bold = True
        else:
            run = paragraph.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

def create_doc_stable(content, ten_bai, lop):
    doc = Document()
    
    # [CẤU HÌNH KHỔ GIẤY A4 VÀ LỀ CHUẨN NĐ30]
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.2
    
    # 1. TIÊU ĐỀ
    head = doc.add_heading(f'KẾ HOẠCH BÀI DẠY: {ten_bai.upper()}', 0)
    head.alignment = 1 
    for run in head.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    p_lop = doc.add_paragraph(f'Lớp: {lop}')
    p_lop.alignment = 1 
    p_lop.runs[0].bold = True
    
    doc.add_paragraph("-" * 60).alignment = 1
    
    # 2. XỬ LÝ NỘI DUNG
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Xóa dấu # đầu dòng
        if line.startswith('#'):
            line = line.replace('#', '').strip()
        
        # [XỬ LÝ BẢNG]
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            if len(table_lines) >= 3: 
                try:
                    valid_rows = [r for r in table_lines if '---' not in r]
                    if valid_rows:
                        cols_count = len(valid_rows[0].split('|')) - 2 
                        if cols_count > 0:
                            table = doc.add_table(rows=len(valid_rows), cols=cols_count)
                            table.style = 'Table Grid'
                            table.autofit = True
                            
                            for r_idx, r_text in enumerate(valid_rows):
                                cells_data = r_text.split('|')[1:-1]
                                for c_idx, cell_text in enumerate(cells_data):
                                    if c_idx < cols_count:
                                        cell = table.cell(r_idx, c_idx)
                                        cell._element.clear_content()
                                        
                                        # Xử lý nội dung ô (Fix lỗi dòng trống như phiên bản trước)
                                        raw_content = cell_text.strip().replace('<br>', '\n').replace('<br/>', '\n')
                                        sub_lines = raw_content.split('\n')
                                        
                                        for sub_line in sub_lines:
                                            sub_line = sub_line.strip()
                                            if not sub_line: continue
                                            
                                            p = cell.add_paragraph()
                                            p.paragraph_format.space_before = Pt(0)
                                            p.paragraph_format.space_after = Pt(2)
                                            p.paragraph_format.line_spacing = 1.1
                                            
                                            if r_idx == 0: # Tiêu đề bảng
                                                p.alignment = 1
                                                run = p.add_run(sub_line.replace('**',''))
                                                run.bold = True
                                                run.font.name = 'Times New Roman'
                                                run.font.size = Pt(14)
                                            else:
                                                add_formatted_text(p, sub_line)
                except: pass
            continue 
            
        if not line:
            i += 1
            continue
            
        # Xử lý Heading (I. II. III.)
        if re.match(r'^(I\.|II\.|III\.|IV\.|V\.)', line) or \
           (re.match(r'^\d+\.', line) and len(line) < 50): 
            clean = line.replace('**', '').strip()
            p = doc.add_paragraph(clean)
            p.runs[0].bold = True
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(14)
        
        # [THAY ĐỔI THEO YÊU CẦU]: Gạch đầu dòng THỦ CÔNG (Không dùng Auto Bullet)
        elif line.startswith('- ') or line.startswith('- '):
            clean = line[2:].strip()
    # Lệnh này khiến Word tự động thụt dòng (Auto Bullet) -> KHÔNG DÙNG NỮA
            p = doc.add_paragraph(style='List Bullet') 
            add_formatted_text(p, clean)
            
        # Đoạn văn thường
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line)
            
        i += 1
        
    return doc

# --- 3. CSS GIAO DIỆN ---
st.markdown("""
<style>
    [data-testid="stAppViewContainer"] { background-color: #f4f6f9; }
    
    .main-header {
        background: linear-gradient(135deg, #004e92 0%, #000428 100%);
        padding: 30px; border-radius: 15px; text-align: center; color: white !important;
        margin-bottom: 30px; box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }
    .main-header h1 { color: white !important; margin: 0; font-family: 'Segoe UI', sans-serif; font-size: 2rem; }
    .main-header p { color: #e0e0e0 !important; margin-top: 10px; font-style: italic; }
    
    .section-header {
        color: #004e92; border-bottom: 2px solid #ddd; padding-bottom: 5px; margin-top: 20px; margin-bottom: 15px; font-weight: bold;
    }
    
    .lesson-plan-paper {
        background-color: white; padding: 40px; border-radius: 5px;
        border: 1px solid #ccc; box-shadow: 0 5px 15px rgba(0,0,0,0.1);
        font-family: 'Times New Roman', Times, serif !important; font-size: 14pt !important;
        line-height: 1.5 !important; color: #000000 !important; text-align: justify;
    }
    .lesson-plan-paper table { width: 100%; border-collapse: collapse; margin: 20px 0; }
    .lesson-plan-paper th, .lesson-plan-paper td { border: 1px solid black; padding: 10px; }
    
    div.stButton > button {
        background: linear-gradient(90deg, #11998e, #38ef7d); color: white !important;
        border: none; padding: 15px 30px; font-weight: bold; border-radius: 10px; width: 100%; margin-top: 10px;
        font-size: 18px; box-shadow: 0 4px 6px rgba(0,0,0,0.2);
    }
    div.stButton > button:hover { transform: translateY(-2px); box-shadow: 0 6px 12px rgba(0,0,0,0.3); }
</style>
""", unsafe_allow_html=True)

# --- 4. GIAO DIỆN CHÍNH ---
st.markdown("""
<div class="main-header">
    <h1>📘 TRỢ LÝ SOẠN GIÁO ÁN TỰ ĐỘNG (NLS)</h1>
    <p>Tác giả: Trần Văn Thọ - TRƯỜNG PTDTBT THCS MƯỜNG LẠN - ĐT: 097 3989 800</p>
</div>
""", unsafe_allow_html=True)

if "GEMINI_API_KEY" in st.secrets:
    api_key = st.secrets["GEMINI_API_KEY"]
else:
    with st.sidebar:
        st.header("🔐 Cấu hình")
        api_key = st.text_input("Nhập API Key:", type="password")

if api_key:
    genai.configure(api_key=api_key)

# --- BỐ TRÍ DỌC (VERTICAL LAYOUT) ---

# 1. TÀI LIỆU
st.markdown('<div class="section-header">📂 1. TÀI LIỆU NGUỒN</div>', unsafe_allow_html=True)

has_framework = False
if os.path.exists(FILE_KHUNG_NANG_LUC):
    st.success(f"✅ Đã tự động tích hợp: {FILE_KHUNG_NANG_LUC}")
    has_framework = True
else:
    st.info(f"ℹ️ Chưa có file '{FILE_KHUNG_NANG_LUC}'. Thầy có thể upload để dùng tính năng Năng lực số.")

uploaded_files = st.file_uploader(
    "Tải Ảnh/PDF bài dạy (Kéo thả vào đây):", 
    type=["jpg", "png", "pdf"], accept_multiple_files=True
)

if uploaded_files:
    st.caption("👁️ Xem trước tài liệu:")
    cols = st.columns(3)
    for i, f in enumerate(uploaded_files):
        if f.type in ["image/jpeg", "image/png"]: 
            with cols[i%3]: st.image(f, caption=f.name)
        else: 
            with cols[i%3]: st.info(f"📄 {f.name}")

# 2. THÔNG TIN
st.markdown('<div class="section-header">📝 2. THÔNG TIN BÀI DẠY</div>', unsafe_allow_html=True)

c1, c2 = st.columns(2)
with c1: lop = st.text_input("📚 Lớp:", "Lớp 6")
with c2: ten_bai = st.text_input("📌 Tên bài học:", placeholder="Ví dụ: Học hát bài...")

noidung_bosung = st.text_area("✍️ Ghi chú thêm (nội dung/kiến thức):", height=100)
yeu_cau_them = st.text_input("💡 Yêu cầu đặc biệt:", placeholder="Ví dụ: Tích hợp trò chơi khởi động...")

# 3. NÚT XỬ LÝ
st.markdown("<br>", unsafe_allow_html=True)
if st.button("🚀 SOẠN GIÁO ÁN NGAY"):
    if not api_key: st.toast("Thiếu API Key!", icon="❌")
    elif not uploaded_files and not noidung_bosung and not has_framework: st.toast("Thiếu tài liệu!", icon="⚠️")
    else:
        try:
            with st.spinner('AI đang soạn giáo án (Times New Roman 14pt, A4, Căn lề chuẩn)...'):
                model = genai.GenerativeModel('gemini-2.5-flash-lite-preview-09-2025')
                
                # --- PROMPT CHI TIẾT CỦA THẦY (BẢN GỐC ĐẦY ĐỦ) ---
                prompt_instruction = f"""
                Đóng vai là một Giáo viên THCS giỏi, am hiểu chương trình GDPT 2018.
                Nhiệm vụ: Soạn Kế hoạch bài dạy (Giáo án) cho bài: "{ten_bai}" - {lop}.

                DỮ LIỆU ĐẦU VÀO:
                - (Nếu có) File PDF Khung năng lực số đính kèm: Hãy dùng để đối chiếu nội dung bài học và đưa vào mục Năng lực số.
                - Các tài liệu hình ảnh/PDF thầy cô tải lên: Phân tích để lấy nội dung kiến thức bài học.
                - Ghi chú bổ sung: "{noidung_bosung}".

                YÊU CẦU LUÔN LUÔN TUÂN THỦ CẤU TRÚC (CÔNG VĂN 2345):
                I. Yêu cầu cần đạt: Trong phần này lại chia thành các phần sau: 
                1. Học sinh thực hiện được, 
                2. Học sinh vận dụng được, 
                3. Phát triển năng lực (bao gồm năng lực đặc thù, năng lực chung, phát triển năng lực số), 
                4. Phát triển phẩm chất.
                * Nội dung tích hợp (VD: Học thông qua chơi, Công dân số,...)
                    - Lưu ý: Thêm phát triển năng lực số trong mục phát triển năng lực (Dựa vào file Khung năng lực nếu có).
                    - Nội dung tích hợp Học thông qua chơi trong Yêu cầu cần đạt cần cụ thể chi tiết hơn chút nữa.
                    - Nội dung tích hợp Công dân số cũng cần cụ thể hơn trong yêu cầu cần đạt.

                II. Đồ dùng dạy học
                1. Giáo viên
                2. Học sinh

                III. Tiến trình dạy học
                [QUAN TRỌNG] PHẦN NÀY PHẢI TRÌNH BÀY DƯỚI DẠNG BẢNG (MARKDOWN TABLE) 2 CỘT. 
                TÊN CÁC HOẠT ĐỘNG PHẢI NẰM Ở CỘT 1 (HOẠT ĐỘNG GIÁO VIÊN).

                | HOẠT ĐỘNG CỦA GIÁO VIÊN | HOẠT ĐỘNG CỦA HỌC SINH |
                |---|---|
                | **1. Hoạt động 1 - Khởi động:**<br>- GV tổ chức... | - HS tham gia... |
                | **2. Hoạt động 2 - Hình thành kiến thức mới:**<br>- GV hướng dẫn... | - HS quan sát... |
                | **3. Hoạt động 3 - Thực hành - luyện tập:**<br>- GV yêu cầu... | - HS thực hiện... |
                | **4. Hoạt động 4 - Vận dụng:**<br>- GV gợi mở... | - HS chia sẻ... |

                YÊU CẦU CHI TIẾT CHO TIẾN TRÌNH DẠY HỌC:
                - Cần chi tiết cụ thể (đặc biệt là Hoạt động của học sinh).
                - Các ý trong tiến trình dạy học được bắt đầu bằng dấu gạch đầu dòng (-).
                - Tích hợp Học thông qua chơi vào 1 số hoạt động phù hợp.
                - Riêng các trò chơi trong tiến trình dạy học cần TRÌNH BÀY RÕ LUẬT CHƠI.
                - Không cần ghi "Mục tiêu HTQC".
                - Tiết học chỉ có 45 phút, hãy điều chỉnh lượng kiến thức và hoạt động hợp lý.
                - Không kèm chú thích nguồn trong bài soạn.
                - Tuyệt đối chỉ bao gồm 4 Hoạt động, không phát sinh thêm.
                - LUÔN LUÔN TUÂN THỦ THEO NHỮNG YÊU CẦU TRÊN

                IV. Điều chỉnh sau tiết dạy

                Lưu ý chung: Bám sát nội dung trong Sách giáo khoa và sách giáo viên (từ tài liệu đính kèm) để đưa nội dung vào bài soạn cho chính xác. KHÔNG dùng ký tự # ở đầu dòng.

                LƯU Ý QUAN TRỌNG TỪ NGƯỜI DÙNG: {yeu_cau_them}
                """

                input_data = [prompt_instruction]
                temp_paths = []
                
                if has_framework: input_data.append(genai.upload_file(FILE_KHUNG_NANG_LUC))
                
                if uploaded_files:
                    for f in uploaded_files:
                        if f.type == "application/pdf":
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                                tmp.write(f.getvalue())
                                temp_paths.append(tmp.name)
                            input_data.append(genai.upload_file(tmp.name))
                        else:
                            input_data.append(Image.open(f))
                
                if noidung_bosung: input_data.append(noidung_bosung)
                
                response = model.generate_content(input_data)
                
                st.markdown("### 📄 KẾT QUẢ BÀI SOẠN:")
                st.markdown(f'<div class="lesson-plan-paper">{response.text}</div>', unsafe_allow_html=True)
                
                doc = create_doc_stable(response.text, ten_bai, lop)
                buf = io.BytesIO()
                doc.save(buf)
                buf.seek(0)
                
                st.download_button(
                    label="⬇️ TẢI FILE WORD CHUẨN A4",
                    data=buf,
                    file_name=f"GiaoAn_{ten_bai}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )
                
                for p in temp_paths: os.remove(p)

        except Exception as e:
            st.error(f"Có lỗi xảy ra: {e}")

# --- CHÂN TRANG ---
st.markdown("---")
st.markdown("<div style='text-align: center; color: #666;'>© 2026 – Trần Văn Thọ - TRƯỜNG PTDTBT THCS MƯỜNG LẠN - ĐT: 097 3989 800</div>", unsafe_allow_html=True)
